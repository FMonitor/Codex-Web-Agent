import docx

class DocxEditor:
    """用于编辑现有 .docx 文档的类。"""

    def __init__(self, file_path):
        self.file_path = file_path
        self.doc = None

    def load(self):
        """加载文档。"""
        try:
            self.doc = docx.Document(self.file_path)
        except Exception as e:
            raise Exception(f"无法加载文档: {str(e)}")

    def replace_text(self, old_text, new_text):
        """替换文档中的特定文本。"""
        if not self.doc:
            self.load()
        
        for para in self.doc.paragraphs:
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)

    def add_paragraph_to_end(self, text, style=None):
        """在文档末尾添加一个段落。"""
        if not self.doc:
            self.load()
        self.doc.add_paragraph(text, style=style)

    def save(self, file_path=None):
        """保存修改后的文档。"""
        if not self.doc:
            self.load()
        
        target_path = file_path if file_path else self.file_path
        try:
            self.doc.save(target_path)
        except Exception as e:
            raise Exception(f"无法保存文档: {str(e)}")
