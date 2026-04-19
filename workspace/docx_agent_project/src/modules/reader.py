import docx

class DocxReader:
    """用于读取 .docx 文档内容的类。"""

    def __init__(self, file_path):
        self.file_path = file_path
        self.doc = None

    def load(self):
        """加载文档。"""
        try:
            self.doc = docx.Document(self.file_path)
        except Exception as e:
            raise Exception(f"无法加载文档: {str(e)}")

    def get_all_text(self):
        """获取文档中所有的文本内容。"""
        if not self.doc:
            self.load()
        return "\n".join([para.text for para in self.doc.paragraphs])

    def get_paragraphs(self):
        """获取所有段落文本列表。"""
        if not self.doc:
            self.load()
        return [para.text for para in self.doc.paragraphs]

    def find_text(self, search_text):
        """查找文本在文档中的段落索引。"""
        if not self.doc:
            self.load()
        for i, para in enumerate(self.doc.paragraphs):
            if search_text in para.text:
                return i
        return -1
