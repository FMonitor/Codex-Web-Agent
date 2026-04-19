import docx

class DocxWriter:
    """用于创建新 .docx 文档的类。"""

    def __init__(self):
        self.doc = docx.Document()

    def add_paragraph(self, text, style=None):
        """添加一个段落。"""
        self.doc.add_paragraph(text, style=style)

    def add_heading(self, text, level=1):
        """添加一个标题。"""
        self.doc.add_heading(text, level=level)

    def save(self, file_path):
        """保存文档。"""
        try:
            self.doc.save(file_path)
        except Exception as e:
            raise Exception(f"无法保存文档: {str(e)}")
