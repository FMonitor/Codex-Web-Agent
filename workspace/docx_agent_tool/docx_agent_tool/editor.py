from docx import Document
from typing import List, Optional, Dict, Any

class DocxEditor:
    """
    A high-level editor for .docx files designed for Agentic workflows.
    """

    def __init__(self, file_path: Optional[str] = None):
        """
        Initialize the editor.
        :param file_path: Path to an existing .docx file. If None, a new document is created.
        """
        if file_path:
            self.file_path = file_path
            self.doc = Document(file_path)
        else:
            self.file_path = None
            self.doc = Document()

    def save(self, file_path: Optional[str] = None):
        """
        Save the document.
        :param file_path: Path to save the document. If None, uses the existing file_path.
        """
        target_path = file_path if file_path else self.file_path
        if not target_path:
            raise ValueError("No file path specified for saving.")
        self.doc.save(target_path)
        self.file_path = target_path

    def add_paragraph(self, text: str, style: Optional[str] = None) -> str:
        """
        Add a paragraph to the document.
        :param text: The text to add.
        :param style: Optional style name.
        :return: The text of the paragraph.
        """
        p = self.doc.add_paragraph(text, style=style)
        return text

    def add_heading(self, text: str, level: int = 1) -> str:
        """
        Add a heading to the document.
        :param text: The text for the heading.
        :param level: The heading level (1-9).
        :return: The text of the heading.
        """
        p = self.doc.add_heading(text, level=level)
        return text

    def add_table(self, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
        """
        Add a table to the document.
        :param rows: Number of rows.
        :param cols: Number of columns.
        :param data: Initial data for the table.
        :return: A summary description of the table.
        """
        table = self.doc.add_table(rows=rows, cols=cols)
        if data:
            for r, row_data in enumerate(data):
                for c, cell_text in enumerate(row_data):
                    if r < rows and c < cols:
                        table.cell(r, c).text = cell_text
        return f"Table with {rows} rows and {cols} columns added."

    def get_text(self) -> str:
        """
        Read all text from the document.
        :return: The full text content.
        """
        full_text = []
        # Note: docx doesn't include table text in paragraphs by default
        for para in self.doc.paragraphs:
            full_text.append(para.text)
        
        # To include table text, we need to iterate through tables
        for table in self.doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
                
        return "\n".join(full_text)

    def list_paragraphs(self) -> List[Dict[str, Any]]:
        """
        List all paragraphs with their index and text.
        :return: A list of dictionaries containing paragraph info.
        """
        return [{"index": i, "text": p.text} for i, p in enumerate(self.doc.paragraphs)]
