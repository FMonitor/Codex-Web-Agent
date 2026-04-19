import docx

class DocxReader:
    """A class to read content from a .docx file."""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def read_text(self) -> str:
        """Reads all text from the docx file."""
        try:
            doc = docx.Document(self.file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            return f"Error reading file: {e}"

    def read_paragraphs(self) -> list[str]:
        """Reads all paragraphs from the docx file as a list of strings."""
        try:
            doc = docx.Document(self.file_path)
            return [para.text for para in doc.paragraphs]
        except Exception as e:
            return [f"Error reading file: {e}"]
